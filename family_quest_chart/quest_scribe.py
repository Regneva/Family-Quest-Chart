"""Class for Character"""
# 1. Standard python libraries
import os

# 2. Third party modules
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR
from docx.shared import Inches

# 3. Family modules

# 4. Local modules

__copyright__ = "(C) Copyright Aaron Jones 2022"
__license__ = "All rights reserved"


class QuestScribe:
    """Chronicle of fellowships and their quests for a day."""
    def __init__(self, directory):
        """Initialize the character class"""
        self.basename = 'AssignedQuests.docx'
        self.filename = os.path.join(directory, self.basename)
        
    def write_highlighted_header(self, doc, first_name, header_text, level=3, highlight_color=WD_COLOR.YELLOW, font_color = RGBColor(0, 0, 0,)):  # RGBColor(255, 255, 0)):
      heading = doc.add_heading(level=level)
      heading.add_run(first_name).font.highlight_color = highlight_color  # Yellow highlighting
      heading.runs[0].font.color.rgb = font_color
      heading.add_run(header_text)
        
    def create_word_document(self, assigned_quests, remaining_quests):
        # Create a new Word document
        doc = Document()
        current_directory = os.path.dirname(os.path.abspath(__file__))
        chk_file = os.path.join(current_directory, 'checkbox.png')

        pageIndex = 1
        for date in assigned_quests:
          if pageIndex != 1:
            doc.add_page_break()
          fellowship_count = 1    
          # Add a title
          doc.add_heading(date.strftime('%A, %B %d, %Y'), level=0)

          for fellowship in assigned_quests[date]:
              # if fellowship_count > 1:
              #   doc.add_paragraph('')
              fellowship_name = f'Fellowship {fellowship_count};'
              if fellowship.mentor:
                fellowship_name += f' {fellowship.mentor.first_name[0]}'
              if fellowship.mentor and fellowship.mentee:
                fellowship_name += ' &'
              if fellowship.mentee:
                fellowship_name += f' {fellowship.mentee.first_name[0]}'
              fellowship_name += ': Quests Complete ______________________'
              doc.add_heading(fellowship_name, level=1)
              fellowship_count += 1
              
              if fellowship.mentor:
                first_name = fellowship.mentor.first_name
                mentor_name = f': {fellowship.mentor.personal_quests}'
                # doc.add_heading(mentor_name, level=2)
                self.write_highlighted_header(doc, first_name, mentor_name, 3, fellowship.mentor.highlight_color, fellowship.mentor.font_color)
                
                quests = {}
                quests = fellowship.mentor_quest_schedule
                for schedule in quests:
                  for quest in quests[schedule]:
                    font_color = quest.font_color
                    para = doc.add_paragraph()
                    run1 = para.add_run()
                    run1.add_picture(chk_file, width=Inches(0.15), height=Inches(0.17))
                    run2 = para.add_run()
                    run2.add_picture(chk_file, width=Inches(0.15), height=Inches(0.17))
                    para.add_run(f' {schedule}: {quest.name}; {quest.description}').font.color.rgb = font_color
                  
              if fellowship.mentee:
                first_name = fellowship.mentee.first_name
                mentee_name = f': {fellowship.mentee.personal_quests}'
                # doc.add_heading(mentee_name, level=3)
                self.write_highlighted_header(doc, first_name, mentee_name, 3, fellowship.mentee.highlight_color, fellowship.mentee.font_color)
                quests = fellowship.mentee_quest_schedule
                for schedule in quests:
                  for quest in quests[schedule]:
                    font_color = quest.font_color
                    para = doc.add_paragraph()
                    run1 = para.add_run()
                    run1.add_picture(chk_file, width=Inches(0.15), height=Inches(0.17))
                    run2 = para.add_run()
                    run2.add_picture(chk_file, width=Inches(0.15), height=Inches(0.17))
                    para.add_run(f' {schedule}: {quest.name}; {quest.description}').font.color.rgb = font_color
                
          doc.add_paragraph('')
          doc.add_heading('Remaining Quests (for reward or discipline)', level=2)
          for quest_index in remaining_quests[date]:
            for schedule in remaining_quests[date][quest_index]:
              quest_count = 0
              quest_line = schedule + ': '
              for quest in remaining_quests[date][quest_index][schedule]:
                if quest_count > 0:
                  quest_line += ", "
                quest_line += quest.name
                quest_count += 1
              if quest_count > 0:
                doc.add_paragraph(quest_line)
            
          pageIndex += 1

        # Save the document with a specified filename
        doc.save(self.filename)
        print(f"Word document '{self.filename}' created successfully.")

